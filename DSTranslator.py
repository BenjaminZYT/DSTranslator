import streamlit as st
import fitz # PyMuPDF for PDF extraction and simple PDF creation
from docx import Document # python-docx for .docx handling
import io
from openai import OpenAI
from langchain.text_splitter import RecursiveCharacterTextSplitter
import os
import re

# --- Configuration ---
DEEPSEEK_BASE_URL = "https://api.deepseek.com"

# --- Functions ---

def extract_text_from_pdf(uploaded_file):
    """Extracts text page by page from a PDF file."""
    doc = fitz.open(stream=uploaded_file.read(), filetype="pdf")
    text_content = []
    for page_num in range(doc.page_count):
        page = doc.load_page(page_num)
        text_content.append(page.get_text())
    doc.close()
    return "\n\n".join(text_content)

def extract_text_from_docx(uploaded_file):
    """Extracts text from a DOCX file, attempting to preserve some structure (e.g., paragraphs, tables)."""
    document = Document(uploaded_file)
    full_text = []
    for element in document.element.body:
        if hasattr(element, 'tag') and element.tag.endswith('body}p'): # Paragraph
            full_text.append(element.text)
        elif hasattr(element, 'tag') and element.tag.endswith('body}tbl'): # Table
            table_text = []
            for row in element.xpath('.//w:tr', namespaces=document.element.nsmap):
                row_cells = []
                for cell in row.xpath('.//w:tc', namespaces=document.element.nsmap):
                    row_cells.append(cell.text)
                # Format as simple Markdown table row for LLM
                table_text.append("| " + " | ".join([c if c else '' for c in row_cells]) + " |")
            if table_text:
                # Add Markdown table header separator
                if len(table_text) > 0:
                    num_cols = len(table_text[0].split('|')) - 2 # -2 for leading/trailing pipes
                    table_text.insert(1, "|" + "---|" * num_cols)
                full_text.append("\n".join(table_text))
    return "\n\n".join(full_text)

def check_deepseek_language_capability(api_key, target_language):
    """
    Asks DeepSeek if it can translate to the specified language.
    This is an inferential check, not a direct API capability list.
    """
    client = OpenAI(api_key=api_key, base_url=DEEPSEEK_BASE_URL)
    try:
        response = client.chat.completions.create(
            model="deepseek-chat",
            messages=[
                {"role": "system", "content": "You are a helpful AI assistant. Your task is to confirm if you can effectively translate text into a specified language. Respond concisely with 'Yes' or 'No', followed by a brief confirmation or reason."},
                {"role": "user", "content": f"Can you translate text into {target_language}?"}
            ],
            temperature=0.1,
            max_tokens=50
        )
        response_text = response.choices[0].message.content.strip().lower()
        if response_text.startswith("yes"):
            return True, "Alright! We can translate to that language."
        else:
            return False, "We don't have this language capability for now, or it may not be fully supported. Please choose another language or clarify."
    except Exception as e:
        st.error(f"Error checking language capability: {e}. Please ensure your API key is correct and valid.")
        return False, "An error occurred while checking language capability."

def translate_text_with_deepseek(text_chunks, target_language):
    """Translates a list of text chunks using DeepSeek API."""
    translated_chunks = []
    api_key = st.session_state.deepseek_api_key
    if not api_key:
        st.error("DeepSeek API Key is not set. Please enter your key and connect.")
        return []

    client = OpenAI(api_key=api_key, base_url=DEEPSEEK_BASE_URL)

    for i, chunk in enumerate(text_chunks):
        with st.spinner(f"Translating chunk {i+1}/{len(text_chunks)}..."):
            try:
                system_message_content = (
                    f"You are a helpful and meticulous translation assistant. "
                    f"Translate the following text into {target_language}. "
                    f"If the text contains tabular data, preserve the table structure in the translation. "
                    f"Represent tables strictly using Markdown table format (e.g., with pipes `|` and hyphens `-`)."
                    f"Maintain original paragraph breaks for non-table text."
                )
                response = client.chat.completions.create(
                    model="deepseek-chat",
                    messages=[
                        {"role": "system", "content": system_message_content},
                        {"role": "user", "content": chunk}
                    ],
                    temperature=0.7,
                    max_tokens=4096
                )
                translated_text = response.choices[0].message.content
                translated_chunks.append(translated_text)
            except Exception as e:
                st.error(f"Error translating chunk {i+1}: {e}. Make sure your API key is correct and valid.")
                translated_chunks.append(f"[Translation Error for chunk {i+1}]: {chunk}")
    return translated_chunks

def create_docx_from_markdown(markdown_text):
    """Creates a DOCX file from Markdown text, attempting to format tables."""
    document = Document()
    lines = markdown_text.split('\n')
    in_table = False
    current_table_rows = []

    for line in lines:
        stripped_line = line.strip()

        # Check for start/end of a Markdown table (simplified detection)
        if stripped_line.startswith('|') and '|' in stripped_line and '-' in stripped_line and not in_table:
            # Likely start of a table or header separator
            in_table = True
            current_table_rows = []
            if stripped_line.startswith('|-'): # This is a separator line
                continue # Skip it for direct parsing as we handle headers from first row
            else: # This is a potential header row
                header_cells = [cell.strip() for cell in stripped_line.split('|') if cell.strip()]
                current_table_rows.append(header_cells)

        elif in_table and stripped_line.startswith('|') and '|' in stripped_line:
            # Another row in the table
            row_cells = [cell.strip() for cell in stripped_line.split('|') if cell.strip()]
            current_table_rows.append(row_cells)

        elif in_table and not stripped_line: # Blank line after table (or end of document)
            # End of table, process it
            if current_table_rows:
                # Determine max columns to handle inconsistent rows, if any
                num_cols = max(len(row) for row in current_table_rows) if current_table_rows else 0
                if num_cols > 0:
                    table = document.add_table(rows=len(current_table_rows), cols=num_cols)
                    table.autofit = True # Auto-adjust column widths

                    for r_idx, row_data in enumerate(current_table_rows):
                        for c_idx, cell_text in enumerate(row_data):
                            if c_idx < num_cols: # Ensure we don't go out of bounds
                                table.cell(r_idx, c_idx).text = cell_text
            in_table = False
            current_table_rows = []
            if stripped_line: # Add the blank line or next text as a paragraph if not empty
                 document.add_paragraph(line) # Add it as a paragraph if it's actual text after table
        elif not in_table: # Not in a table, add as a paragraph
            document.add_paragraph(line)
        # Handle cases where table definition might be partial or malformed
        # If the line looks like a table row but we're not 'in_table' yet, might be a malformed table,
        # so just add it as a paragraph
        elif stripped_line.startswith('|') and '|' in stripped_line and not in_table:
            document.add_paragraph(line)


    # Handle case where the document ends with a table
    if in_table and current_table_rows:
        num_cols = max(len(row) for row in current_table_rows) if current_table_rows else 0
        if num_cols > 0:
            table = document.add_table(rows=len(current_table_rows), cols=num_cols)
            table.autofit = True
            for r_idx, row_data in enumerate(current_table_rows):
                for c_idx, cell_text in enumerate(row_data):
                    if c_idx < num_cols:
                        table.cell(r_idx, c_idx).text = cell_text


    bio = io.BytesIO()
    document.save(bio)
    return bio.getvalue()

def create_pdf_from_text(text_content):
    """
    Creates a basic PDF from plain text content using PyMuPDF.
    NOTE: This will NOT render Markdown tables as actual PDF tables.
    They will appear as plain text (e.g., | Cell 1 | Cell 2 |).
    Complex layouts and non-Latin characters might require font embedding.
    """
    doc = fitz.open()
    page = doc.new_page()
    
    # Define text settings
    font_name = "helv" # Helvetica, a common font in PyMuPDF
    font_size = 10
    line_height = 12
    margin = 50
    
    # Get page dimensions
    page_width = page.rect.width
    page_height = page.rect.height
    
    # Calculate available text area
    text_area_width = page_width - 2 * margin
    cursor_y = page_height - margin # Start from top margin

    # A simple regex to detect potential markdown table lines
    # This is for display purposes only, not for rendering a true table.
    markdown_table_line_regex = re.compile(r'^\s*\|.*\|')

    for line in text_content.split('\n'):
        # If the line is empty, just move to the next line
        if not line.strip():
            cursor_y -= line_height
            if cursor_y < margin: # New page if content goes beyond bottom margin
                page = doc.new_page()
                cursor_y = page_height - margin
            continue
            
        # Attempt to draw the line
        try:
            # Insert text, handling wrapping.
            # Using insert_textbox is better for wrapping than insert_text
            text_rect = fitz.Rect(margin, cursor_y - line_height, page_width - margin, cursor_y)
            
            # Check if it looks like a markdown table line
            if markdown_table_line_regex.match(line):
                # For table lines, we might want slightly different styling or just raw text
                # We'll just draw it as regular text for simplicity, without table structure
                pass

            # This loop manually handles line breaks and pagination
            # A more robust solution might use text_as_markdown property or a more advanced library
            text_instance = page.insert_textbox(
                text_rect,
                line,
                fontname=font_name,
                fontsize=font_size,
                align=fitz.TEXT_ALIGN_LEFT,
                render_mode=0, # fill mode
                # For non-Latin languages, you *must* embed fonts.
                # Example for Chinese: fontfile="path/to/your/chinese_font.ttf"
            )
            
            # Move cursor down based on actual text height (if it wrapped)
            cursor_y = text_instance.y1 # The bottom of the inserted text box
            
            # If text flowed past the bottom margin, add new page
            if cursor_y < margin:
                page = doc.new_page()
                cursor_y = page_height - margin # Reset cursor for new page
                
        except Exception as e:
            st.warning(f"Could not render line in PDF: '{line[:50]}...' - Error: {e}")
            # Fallback to simple new line
            cursor_y -= line_height

    pdf_bytes = doc.tobytes()
    doc.close()
    return pdf_bytes


# --- Streamlit App Layout ---
st.set_page_config(page_title="DeepSeek Document Translator", layout="centered")

st.title("📄 DeepSeek Document Translator")
st.markdown("Enter your API key, upload a document, specify your target language, and choose your output format!")

# --- API Key Input ---
st.header("1. DeepSeek API Key")
api_key_input = st.text_input(
    "Enter your DeepSeek API Key:",
    type="password",
    key="api_key_input"
)

if 'deepseek_api_key' not in st.session_state:
    st.session_state.deepseek_api_key = None
if 'language_capability_checked' not in st.session_state:
    st.session_state.language_capability_checked = False
if 'can_translate' not in st.session_state:
    st.session_state.can_translate = False

if st.button("Connect to DeepSeek"):
    if api_key_input:
        st.session_state.deepseek_api_key = api_key_input
        st.success("API Key set! You can now proceed to upload your document.")
    else:
        st.error("Please enter your DeepSeek API Key.")

# --- Document Upload and Options ---
if st.session_state.deepseek_api_key:
    st.header("2. Upload Document & Set Translation Options")
    uploaded_file = st.file_uploader(
        "Upload your document (PDF or DOCX)",
        type=["pdf", "docx"]
    )

    target_language_input = st.text_input(
        "Enter Target Language (e.g., 'Japanese', 'Simplified Chinese'):",
        key="target_language_input"
    )

    if target_language_input and st.button("Check Language Capability"):
        can_translate, message = check_deepseek_language_capability(st.session_state.deepseek_api_key, target_language_input)
        st.session_state.can_translate = can_translate
        st.session_state.language_capability_checked = True
        if can_translate:
            st.success(message)
        else:
            st.warning(message)
    elif not target_language_input:
        st.session_state.language_capability_checked = False

    output_format = st.selectbox(
        "Desired Output File Format:",
        ["PDF", "DOCX"],
        index=1 # Default to DOCX as it will handle tables better
    )

    if uploaded_file and st.session_state.language_capability_checked and st.session_state.can_translate:
        if st.button("Translate Document"):
            original_file_extension = uploaded_file.name.split('.')[-1].lower()
            full_text_content = ""

            with st.spinner("Extracting text from document..."):
                if original_file_extension == "pdf":
                    full_text_content = extract_text_from_pdf(uploaded_file)
                elif original_file_extension == "docx":
                    full_text_content = extract_text_from_docx(uploaded_file)
                else:
                    st.error("Unsupported file type. Please upload a PDF or DOCX.")
                    st.stop()

            if not full_text_content:
                st.warning("Could not extract text from the document. It might be empty or corrupted.")
                st.stop()

            text_splitter = RecursiveCharacterTextSplitter(
                chunk_size=3000,
                chunk_overlap=200,
                length_function=len,
            )
            text_chunks = text_splitter.split_text(full_text_content)
            st.info(f"Extracted text and split into {len(text_chunks)} chunks for translation.")

            translated_chunks = translate_text_with_deepseek(text_chunks, target_language_input)

            if translated_chunks:
                translated_content = "\n\n".join(translated_chunks)

                st.subheader("3. Download Translated Document")
                translated_file_name = f"translated_{os.path.splitext(uploaded_file.name)[0]}.{output_format.lower()}"

                translated_bytes = None
                if output_format == "DOCX":
                    with st.spinner("Creating translated Word document..."):
                        translated_bytes = create_docx_from_markdown(translated_content)
                elif output_format == "PDF":
                    with st.spinner("Creating translated PDF... (Note: Tables will appear as plain text, not formatted tables.)"):
                        translated_bytes = create_pdf_from_text(translated_content)

                if translated_bytes:
                    st.success("Translation and conversion complete!")
                    st.download_button(
                        label=f"Download Translated {output_format}",
                        data=translated_bytes,
                        file_name=translated_file_name,
                        mime=f"application/{output_format.lower()}" if output_format == "PDF" else "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
                else:
                    st.error("Failed to generate the translated document.")
            else:
                st.error("Translation failed. Please check your API key and chosen language, then try again.")
    elif uploaded_file and not st.session_state.language_capability_checked:
        st.info("Please enter the target language and click 'Check Language Capability' first.")
    elif st.session_state.language_capability_checked and not st.session_state.can_translate:
        st.warning("Cannot proceed with translation as the chosen language is not supported or encountered an error.")
    elif not uploaded_file:
        st.info("Please upload a document file to begin.")
else:
    st.info("Please enter your DeepSeek API Key and click 'Connect' to proceed.")